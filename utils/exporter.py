"""
Document Exporter Module

This module handles exporting generated content to various file formats
including PDF, DOCX, TXT, and Markdown using industry-standard libraries.
"""
import io
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

def export_txt(content: str) -> bytes:
    """Exports content as a TXT file byte string."""
    return content.encode("utf-8")

def export_markdown(content: str) -> bytes:
    """Exports content as a Markdown file byte string."""
    return content.encode("utf-8")

def export_docx(content: str) -> bytes:
    """Exports content as a DOCX file byte string."""
    doc = Document()
    
    # Simple text to DOCX mapping
    for line in content.split('\n'):
        clean_line = line.strip()
        if not clean_line:
            continue
            
        if clean_line.startswith('# '):
            doc.add_heading(clean_line.replace('# ', ''), level=1)
        elif clean_line.startswith('## '):
            doc.add_heading(clean_line.replace('## ', ''), level=2)
        elif clean_line.startswith('### '):
            doc.add_heading(clean_line.replace('### ', ''), level=3)
        elif clean_line.startswith('- ') or clean_line.startswith('* '):
            doc.add_paragraph(clean_line[2:], style='List Bullet')
        else:
            doc.add_paragraph(clean_line)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def export_pdf(content: str) -> bytes:
    """Exports content as a PDF file using ReportLab."""
    pdf_io = io.BytesIO()
    doc = SimpleDocTemplate(pdf_io, pagesize=letter)
    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]
    heading1_style = styles["Heading1"]
    heading2_style = styles["Heading2"]
    
    flowables = []
    
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            flowables.append(Spacer(1, 12))
            continue
            
        # Strip bold/italic markdown for ReportLab basic parsing or escape XML chars
        safe_line = line.replace('<', '&lt;').replace('>', '&gt;')
        # Remove common markdown bolding to prevent ReportLab XML errors
        safe_line = re.sub(r'\*\*(.*?)\*\*', r'\1', safe_line)
            
        if safe_line.startswith('# '):
            flowables.append(Paragraph(safe_line.replace('# ', ''), heading1_style))
        elif safe_line.startswith('## '):
            flowables.append(Paragraph(safe_line.replace('## ', ''), heading2_style))
        elif safe_line.startswith('### '):
            flowables.append(Paragraph(safe_line.replace('### ', ''), heading2_style))
        else:
            flowables.append(Paragraph(safe_line, normal_style))
            
    doc.build(flowables)
    pdf_io.seek(0)
    return pdf_io.getvalue()
